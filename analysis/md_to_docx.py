# -*- coding: utf-8 -*-
"""報告書Markdown → Word(docx)変換。表は罫線表、コードブロックは等幅の枠付き段落で出力する。"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC, DST = sys.argv[1], sys.argv[2]
doc = Document()

# 日本語フォント設定
st = doc.styles['Normal']
st.font.name = 'Yu Gothic'; st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
for s in ('Heading 1','Heading 2','Heading 3','Heading 4'):
    h = doc.styles[s]; h.font.name = 'Yu Gothic'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    h.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.7)
    s.top_margin = s.bottom_margin = Inches(0.7)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')
def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        else:
            r = par.add_run(tok)
        r.font.name = r.font.name or 'Yu Gothic'
        try: r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        except Exception: pass

def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    ln = lines[i]
    # コードブロック
    if ln.startswith('```'):
        i += 1; buf = []
        while i < len(lines) and not lines[i].startswith('```'):
            buf.append(lines[i]); i += 1
        i += 1
        tb = doc.add_table(rows=1, cols=1); tb.style = 'Table Grid'
        c = tb.cell(0,0); c.text = ''
        for k, b in enumerate(buf):
            p = c.paragraphs[0] if k == 0 else c.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(b); r.font.name = 'Consolas'; r.font.size = Pt(7.5)
        shade(c, 'F5F6F8')
        doc.add_paragraph()
        continue
    # 表
    if ln.startswith('|') and i+1 < len(lines) and re.match(r'^\|[\s:\-\|]+\|$', lines[i+1].strip()):
        hdr = split_row(ln); i += 2; rows = []
        while i < len(lines) and lines[i].startswith('|'):
            rows.append(split_row(lines[i])); i += 1
        n = len(hdr)
        tb = doc.add_table(rows=1, cols=n); tb.style = 'Table Grid'
        tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(hdr):
            cell = tb.rows[0].cells[j]; cell.text = ''
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            add_runs(p, h if h.startswith('**') else f'**{h}**')
            for r in p.runs: r.font.size = Pt(8.5)
            shade(cell, 'DCE6F1')
        for row in rows:
            cells = tb.add_row().cells
            for j in range(n):
                v = row[j] if j < len(row) else ''
                cells[j].text = ''
                p = cells[j].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
                add_runs(p, v)
                for r in p.runs: r.font.size = Pt(8.5)
        doc.add_paragraph()
        continue
    # 見出し
    m = re.match(r'^(#{1,4})\s+(.*)$', ln)
    if m:
        lvl = len(m.group(1)); txt = m.group(2)
        if lvl == 1:
            p = doc.add_heading('', level=1); add_runs(p, txt)
        else:
            p = doc.add_heading('', level=min(lvl,4)); add_runs(p, txt)
        i += 1; continue
    if ln.strip() == '---':
        doc.add_paragraph('─'*60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue
    if ln.startswith('> '):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, ln[2:])
        for r in p.runs: r.italic = True
        i += 1; continue
    m = re.match(r'^(\s*)[-*]\s+(.*)$', ln)
    if m:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 + 0.25*(len(m.group(1))//2))
        add_runs(p, m.group(2)); i += 1; continue
    m = re.match(r'^(\s*)(\d+)\.\s+(.*)$', ln)
    if m:
        p = doc.add_paragraph(style='List Number'); add_runs(p, m.group(3)); i += 1; continue
    if ln.strip() == '':
        i += 1; continue
    p = doc.add_paragraph(); add_runs(p, ln)
    i += 1

doc.save(DST)
print('saved:', DST)
